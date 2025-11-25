from docx import Document
import json
import re

# Cấu trúc bảng:
# STT | Nội dung câu hỏi | Phương án | Đáp án


def extract_options_from_cell(cell):
    paras = [p.text.strip() for p in cell.paragraphs if p.text.strip()]
    if not paras:
        return {}

    full_text = "\n".join(paras)

    full_text = re.sub(r'\s+', ' ', full_text).strip()

    full_text = re.sub(r'(?<!\n)\s*([A-Da-d])[).:\-]\s*', r'\n\1. ', full_text)

    lines = [line.strip() for line in full_text.split("\n") if line.strip()]

    options = {}
    current_key = None
    current_val = ""

    for line in lines:
        match = re.match(r'^([A-Da-d])[).:\-]?\s+(.*)', line)
        if match:
            if current_key:
                options[current_key] = current_val.strip()

            current_key = match.group(1).lower()  # a/b/c/d
            current_val = match.group(2).strip()
        else:
            if current_key:
                current_val += " " + line.strip()

    if current_key:
        options[current_key] = current_val.strip()

    if len(options) == 4:
        return options


    fallback_options = {}
    labels = ["a", "b", "c", "d"]
    for i, para_text in enumerate(paras[:4]): 
        fallback_options[labels[i]] = para_text

    return fallback_options


def safe_parse_stt(raw_text, default=None):

    text = (raw_text or "").strip()
    if not text:
        return default

    m = re.search(r'\d+', text)
    if not m:
        return default

    try:
        return int(m.group(0))
    except ValueError:
        return default


def normalize_answer_cell(text):

    t = (text or "").strip().upper()
    m = re.search(r'[ABCD]', t)
    if not m:
        return ""
    return m.group(0)


def parse_questions_from_table(docx_path, output_json_path):
    document = Document(docx_path)
    questions = []

    global_stt_counter = 0  # Đếm tổng số câu đã parse (cross-table)

    for table_idx, table in enumerate(document.tables):
        # Bỏ qua dòng tiêu đề, duyệt từ dòng thứ 2
        for row_idx, row in enumerate(table.rows[1:], start=1):
            cells = row.cells

            # Bảo vệ: nếu số cột < 4 thì bỏ qua
            if len(cells) < 4:
                print(f"[Table {table_idx+1} Row {row_idx+1}] Bỏ qua vì số cột < 4")
                continue

            try:
                raw_stt = cells[0].text
                noidung_raw = cells[1].text
                options_cell = cells[2]
                dap_an_raw = cells[3].text

                # STT: cố gắng lấy số trong ô 0, nếu không được → dùng counter
                stt = safe_parse_stt(raw_stt, default=None)
                if stt is None:
                    global_stt_counter += 1
                    stt = global_stt_counter
                else:
                    global_stt_counter = stt

                noidung = noidung_raw.strip()

                phuongan_dict = extract_options_from_cell(options_cell)
                dap_an = normalize_answer_cell(dap_an_raw)

                cau_hoi = {
                    "stt": stt,
                    "noidung": noidung,
                    "a": phuongan_dict.get("a", ""),
                    "b": phuongan_dict.get("b", ""),
                    "c": phuongan_dict.get("c", ""),
                    "d": phuongan_dict.get("d", ""),
                    "dapandung": dap_an
                }

                # Kiểm tra tối thiểu: phải có ít nhất 2 phương án, và có đáp án A–D
                non_empty_opts = [k for k, v in phuongan_dict.items() if v.strip()]
                if len(non_empty_opts) >= 2 and dap_an in ["A", "B", "C", "D"]:
                    questions.append(cau_hoi)
                else:
                    print(
                        f"[Câu {stt}] BỎ: thiếu phương án/đáp án. "
                        f"Options có: {sorted(phuongan_dict.keys())}, đáp án: '{dap_an_raw.strip()}'"
                    )

            except Exception as e:
                id_str = cells[0].text.strip() if cells[0].text.strip() else f"Row {row_idx+1}"
                print(f"Lỗi khi xử lý dòng {id_str}: {e}")
                continue

    with open(output_json_path, "w", encoding="utf-8") as f:
        json.dump(questions, f, ensure_ascii=False, indent=2)

    print(f"\nĐã chuyển {len(questions)} câu hỏi hợp lệ sang {output_json_path}")


if __name__ == "__main__":
    input_docx = "data/cauhoichuong5.docx"
    output_json = "output/questions5.json"
    parse_questions_from_table(input_docx, output_json)